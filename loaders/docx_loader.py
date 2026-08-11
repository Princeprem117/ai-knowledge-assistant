from docx import Document as DocxDocument

from loaders.base_loader import BaseLoader
from models.document import Document


class DocxLoader(BaseLoader):

    def load(self, file_path: str) -> Document:
        """Load text data from a DOCX file."""

        docx = DocxDocument(file_path)

        paragraphs = [
            paragraph.text
            for paragraph in docx.paragraphs
            if paragraph.text.strip()
        ]

        text = "\n".join(paragraphs)

        return Document(
            content=text,
            metadata={
                "source": file_path,
                "type": "docx",
            },
        )